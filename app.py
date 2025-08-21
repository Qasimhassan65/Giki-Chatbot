import os
from pathlib import Path
from typing import List
import re

import pdfplumber
import docx
import gradio as gr

from fastapi import FastAPI
from pydantic import BaseModel
import threading


from langchain.docstore.document import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain_openai import ChatOpenAI


OPENROUTER_API_KEY = os.getenv("GIKI_OPENROUTER_API_KEY")
os.environ["OPENAI_API_KEY"] = OPENROUTER_API_KEY
os.environ["OPENAI_API_BASE"] = "https://openrouter.ai/api/v1"
os.environ["OPENAI_API_HEADERS"] = '{"HTTP-Referer":"https://huggingface.co","X-Title":"GIKI-RAG-bot"}'


# -----------------------------
# Document Processor
# -----------------------------
class GIKIDocumentProcessor:
    def __init__(self, data_folder="data"):
        self.data_folder = data_folder
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=500,
            chunk_overlap=50,
            length_function=len,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
        )

    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        page_text = re.sub(r'\s+', ' ', page_text)
                        text += f"\n[Page {page_num + 1}]\n{page_text}\n"
        except Exception as e:
            print(f"Error extracting from PDF {file_path}: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        text = ""
        try:
            doc = docx.Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting from DOCX {file_path}: {e}")
        return text

    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting from TXT {file_path}: {e}")
            return ""

    def load_documents(self) -> List[Document]:
        documents = []
        data_path = Path(self.data_folder)
        if not data_path.exists():
            print(f"Data folder '{self.data_folder}' not found!")
            return documents

        supported_extensions = {'.pdf', '.docx', '.txt'}
        for file_path in data_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                print(f"Processing: {file_path.name}")
                if file_path.suffix.lower() == '.pdf':
                    text = self.extract_text_from_pdf(str(file_path))
                elif file_path.suffix.lower() == '.docx':
                    text = self.extract_text_from_docx(str(file_path))
                elif file_path.suffix.lower() == '.txt':
                    text = self.extract_text_from_txt(str(file_path))
                else:
                    continue

                if text.strip():
                    chunks = self.text_splitter.split_text(text)
                    for i, chunk in enumerate(chunks):
                        if chunk.strip():
                            documents.append(
                                Document(
                                    page_content=chunk,
                                    metadata={
                                        "source": file_path.name,
                                        "chunk_id": i,
                                        "file_type": file_path.suffix.lower()
                                    }
                                )
                            )
        print(f"✅ Loaded {len(documents)} document chunks")
        return documents


# -----------------------------
# Chatbot
# -----------------------------
class GIKIbot:
    def __init__(self):
        self.qa_chain = None
        self.vectorstore = None
        self.processor = GIKIDocumentProcessor()
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        self.prompt_template = """You are a helpful assistant for GIKI.
Answer questions based on official documents.

Context:
{context}

Question: {question}

Answer:"""
        self.custom_prompt = PromptTemplate(
            template=self.prompt_template,
            input_variables=["context", "question"]
        )

    def initialize_system(self):
        try:
            print("📂 Loading documents...")
            documents = self.processor.load_documents()
            if not documents:
                return "❌ No documents found in 'data' folder."

            print("🔍 Creating embeddings...")
            self.vectorstore = FAISS.from_documents(documents, self.embeddings)

            llm = ChatOpenAI(
                model="openai/gpt-3.5-turbo",
                base_url="https://openrouter.ai/api/v1",
                api_key=os.getenv("GIKI_OPENROUTER_API_KEY"),
                default_headers={"HTTP-Referer": "https://huggingface.co", "X-Title": "GIKI-RAG-bot"},
                temperature=0.1
            )

            self.qa_chain = RetrievalQA.from_chain_type(
                llm=llm,
                chain_type="stuff",
                retriever=self.vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 5}),
                return_source_documents=True,
                chain_type_kwargs={"prompt": self.custom_prompt}
            )
            return "✅ System ready! Ask your questions."
        except Exception as e:
            return f"❌ Error initializing system: {str(e)}"

    def ask_question(self, question: str) -> str:
        if not self.qa_chain:
            return "⚠️ System not initialized yet."
        if not question.strip():
            return "⚠️ Please enter a valid question."
        try:
            response = self.qa_chain({"query": question})
            answer = response["result"]
            sources = set(f"📄 {doc.metadata['source']}" for doc in response["source_documents"])
            source_text = "\n\nSources:\n" + "\n".join(sources) if sources else ""
            return f"{answer}\n\n{source_text}"
        except Exception as e:
            return f"❌ Error: {str(e)}"


# -----------------------------
# Gradio Interface
# -----------------------------
bot = GIKIbot()
init_message = bot.initialize_system()

def chat_fn(user_input):
    return bot.ask_question(user_input)

with gr.Blocks() as demo:
    gr.Markdown("# GIKI-RAG Chatbot")
    gr.Markdown(init_message)

    chatbot = gr.Chatbot()
    msg = gr.Textbox(placeholder="Ask a question...")
    clear = gr.Button("Clear Chat")
    state = gr.State([])

    def respond(user_message, history):
        bot_reply = chat_fn(user_message)
        history = history + [(user_message, bot_reply)]
        return history, history

    msg.submit(respond, [msg, state], [chatbot, state])
    clear.click(lambda: [], None, [chatbot, state])

# --- FastAPI part ---
api = FastAPI()

class Question(BaseModel):
    question: str

@api.post("/chat")
def chat_endpoint(q: Question):
    return {"answer": bot.ask_question(q.question)}

if __name__ == "__main__":
    # Start Gradio in a separate thread (optional)
    threading.Thread(target=lambda: demo.launch(server_name="0.0.0.0", server_port=8081)).start()
    # FastAPI runs on port 8080 (or $PORT for Railway)
    import uvicorn
    uvicorn.run(api, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))